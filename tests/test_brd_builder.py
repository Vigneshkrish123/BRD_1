import io
import pytest
from docx import Document
from unittest.mock import patch

from app.core.models import (
    BRDContent, FunctionalRequirement, NonFunctionalRequirement,
    Risk, Likelihood, Impact
)
from app.services.brd_builder import build_brd, BRDBuilderError
from scripts.create_template import create_template
import tempfile
import os


# ── Fixtures ───────────────────────────────────────────────────────────────────

@pytest.fixture
def sample_content() -> BRDContent:
    return BRDContent(
        project_name="Customer Portal v2",
        meeting_date="2024-11-15",
        attendees=["John Smith (PM)", "Jane Doe (BA)", "Bob Lee (Tech Lead)"],
        business_objectives=[
            "Reduce customer support calls by 30% via self-service portal.",
            "Increase customer satisfaction score to 4.5/5 by Q3.",
        ],
        in_scope=[
            "Customer self-service portal web application",
            "Azure AD SSO integration",
            "Salesforce CRM read/write integration",
        ],
        out_of_scope=[
            "Mobile native app (deferred to Phase 2)",
            "IE11 browser support",
        ],
        functional_requirements=[
            FunctionalRequirement(id="FR-001", description="Users must authenticate via Azure AD SSO.", priority="High"),
            FunctionalRequirement(id="FR-002", description="Portal must display customer order history from Salesforce.", priority="High"),
            FunctionalRequirement(id="FR-003", description="Users must be able to raise and track support tickets.", priority="Medium"),
        ],
        non_functional_requirements=[
            NonFunctionalRequirement(id="NFR-001", category="Scalability", description="System must support 10,000 concurrent users."),
            NonFunctionalRequirement(id="NFR-002", category="Availability", description="System uptime must be 99.9% excluding maintenance windows."),
        ],
        assumptions=[
            "Azure AD tenant is already configured and licensed.",
            "Salesforce sandbox environment will be available for testing.",
        ],
        constraints=[
            "Budget is hard-capped at USD 200,000.",
            "Go-live deadline is Q3 end — no extensions.",
        ],
        risks=[
            Risk(
                id="RISK-001",
                description="Salesforce API deprecation may break integration mid-project.",
                likelihood=Likelihood.MEDIUM,
                impact=Impact.HIGH,
                mitigation="Monitor Salesforce release notes; build abstraction layer.",
            ),
            Risk(
                id="RISK-002",
                description="Azure AD tenant misconfiguration could delay SSO setup.",
                likelihood=Likelihood.LOW,
                impact=Impact.MEDIUM,
                mitigation="Engage IT team in week 1 for AD validation.",
            ),
        ],
    )


@pytest.fixture
def temp_template(tmp_path) -> str:
    """Generate a fresh template in a temp directory for each test."""
    template_path = str(tmp_path / "brd_template.docx")
    create_template(template_path)
    return template_path


# ── Tests ──────────────────────────────────────────────────────────────────────

def test_build_brd_returns_bytes(sample_content, temp_template):
    with patch("app.services.brd_builder.settings") as mock_settings:
        mock_settings.brd_template_path = temp_template
        result = build_brd(sample_content)

    assert isinstance(result, bytes)
    assert len(result) > 1000  # Non-trivial document


def test_build_brd_is_valid_docx(sample_content, temp_template):
    with patch("app.services.brd_builder.settings") as mock_settings:
        mock_settings.brd_template_path = temp_template
        result = build_brd(sample_content)

    # Should open without error
    doc = Document(io.BytesIO(result))
    assert doc is not None


def test_build_brd_project_name_in_document(sample_content, temp_template):
    with patch("app.services.brd_builder.settings") as mock_settings:
        mock_settings.brd_template_path = temp_template
        result = build_brd(sample_content)

    doc = Document(io.BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Customer Portal v2" in full_text


def test_build_brd_fr_table_has_correct_rows(sample_content, temp_template):
    with patch("app.services.brd_builder.settings") as mock_settings:
        mock_settings.brd_template_path = temp_template
        result = build_brd(sample_content)

    doc = Document(io.BytesIO(result))
    # Find the FR table: header row + 3 FR rows = 4 total
    fr_table = None
    for table in doc.tables:
        if table.rows[0].cells[0].text == "ID":
            cell_texts = [c.text for c in table.rows[0].cells]
            if "Priority" in cell_texts:
                fr_table = table
                break

    assert fr_table is not None
    assert len(fr_table.rows) == 4  # 1 header + 3 FRs


def test_build_brd_risk_table_has_correct_rows(sample_content, temp_template):
    with patch("app.services.brd_builder.settings") as mock_settings:
        mock_settings.brd_template_path = temp_template
        result = build_brd(sample_content)

    doc = Document(io.BytesIO(result))
    risk_table = None
    for table in doc.tables:
        headers = [c.text for c in table.rows[0].cells]
        if "Mitigation" in headers:
            risk_table = table
            break

    assert risk_table is not None
    assert len(risk_table.rows) == 3  # 1 header + 2 risks


def test_build_brd_missing_template_raises(sample_content):
    with patch("app.services.brd_builder.settings") as mock_settings:
        mock_settings.brd_template_path = "/nonexistent/path/template.docx"
        with pytest.raises(BRDBuilderError, match="template not found"):
            build_brd(sample_content)


def test_build_brd_empty_sections_handled(temp_template):
    """Empty lists should not crash — should write 'None identified'."""
    minimal = BRDContent(
        project_name="Minimal Project",
        meeting_date="Not specified",
        attendees=[],
        business_objectives=[],
        in_scope=[],
        out_of_scope=[],
        functional_requirements=[],
        non_functional_requirements=[],
        assumptions=[],
        constraints=[],
        risks=[],
    )
    with patch("app.services.brd_builder.settings") as mock_settings:
        mock_settings.brd_template_path = temp_template
        result = build_brd(minimal)

    assert isinstance(result, bytes)
    doc = Document(io.BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "None identified" in full_text
