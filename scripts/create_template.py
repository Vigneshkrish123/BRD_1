"""
Creates the base BRD company template (.docx) with all section placeholders.
Run once to generate templates/brd_template.docx

Usage:
    python scripts/create_template.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ── Brand Colors ───────────────────────────────────────────────────────────────
COLOR_PRIMARY    = RGBColor(0x1F, 0x49, 0x7D)   # Deep navy
COLOR_SECONDARY  = RGBColor(0x2E, 0x75, 0xB6)   # Mid blue
COLOR_ACCENT     = RGBColor(0xED, 0x7D, 0x31)   # Orange accent
COLOR_LIGHT_BG   = RGBColor(0xD6, 0xE4, 0xF0)   # Light blue background
COLOR_TABLE_HDR  = RGBColor(0x1F, 0x49, 0x7D)   # Table header navy
COLOR_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TEXT       = RGBColor(0x26, 0x26, 0x26)   # Near black
COLOR_MUTED      = RGBColor(0x59, 0x59, 0x59)   # Muted grey


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get("val", "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz", "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "CCCCCC"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _add_run(para, text: str, bold=False, italic=False,
             size_pt=11, color: RGBColor = None, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _heading(doc, text: str, level: int):
    """Add a styled heading paragraph."""
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Calibri"

    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = COLOR_PRIMARY
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(6)
        # Bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F497D")
        pBdr.append(bottom)
        pPr.append(pBdr)

    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_SECONDARY
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)

    return para


def _placeholder_para(doc, tag: str, hint: str = ""):
    """Add a tagged placeholder paragraph the builder will find and replace."""
    para = doc.add_paragraph()
    run = para.add_run(f"{{{{ {tag} }}}}")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_MUTED
    run.italic = True
    if hint:
        hint_run = para.add_run(f"  # {hint}")
        hint_run.font.size = Pt(9)
        hint_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        hint_run.italic = True
    return para


def _cover_page(doc):
    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(60)
    _add_run(title, "BUSINESS REQUIREMENTS DOCUMENT",
             bold=True, size_pt=22, color=COLOR_PRIMARY, font="Calibri")

    doc.add_paragraph()  # spacer

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sub, "{{ project_name }}", bold=True,
             size_pt=18, color=COLOR_SECONDARY)

    doc.add_paragraph()

    # Meta table
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"

    labels = ["Project Name", "Meeting Date", "Prepared By", "Version"]
    tags   = ["{{ project_name }}", "{{ meeting_date }}", "BRD Agent", "1.0"]

    for i, (label, tag) in enumerate(zip(labels, tags)):
        lc, vc = meta.rows[i].cells
        _set_cell_bg(lc, "1F497D")
        _set_cell_bg(vc, "F5F9FC")
        lc.paragraphs[0].clear()
        _add_run(lc.paragraphs[0], label, bold=True,
                 color=COLOR_WHITE, size_pt=10)
        vc.paragraphs[0].clear()
        _add_run(vc.paragraphs[0], tag,
                 color=COLOR_TEXT, size_pt=10)

    doc.add_page_break()


def _attendees_section(doc):
    _heading(doc, "Meeting Attendees", 2)
    _placeholder_para(doc, "attendees_list",
                      "Comma-separated names/roles from transcript")


def _section_bullets(doc, title: str, tag: str, hint: str):
    _heading(doc, title, 1)
    _placeholder_para(doc, tag, hint)


def _fr_table(doc):
    _heading(doc, "4. Functional Requirements", 1)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    widths = [Cm(2.5), Cm(10.5), Cm(3.0)]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width

    # Header row
    hdr_cells = table.rows[0].cells
    headers = ["ID", "Description", "Priority"]
    for cell, hdr in zip(hdr_cells, headers):
        _set_cell_bg(cell, "1F497D")
        cell.paragraphs[0].clear()
        _add_run(cell.paragraphs[0], hdr,
                 bold=True, color=COLOR_WHITE, size_pt=10)

    # Placeholder row — builder replaces this
    row = table.add_row()
    _set_cell_bg(row.cells[0], "F5F9FC")
    _set_cell_bg(row.cells[1], "F5F9FC")
    _set_cell_bg(row.cells[2], "F5F9FC")
    row.cells[0].paragraphs[0].clear()
    _add_run(row.cells[0].paragraphs[0], "{{ fr_rows }}",
             italic=True, color=COLOR_MUTED, size_pt=9)


def _nfr_table(doc):
    _heading(doc, "5. Non-Functional Requirements", 1)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    widths = [Cm(2.5), Cm(3.5), Cm(10.0)]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width

    hdr_cells = table.rows[0].cells
    for cell, hdr in zip(hdr_cells, ["ID", "Category", "Description"]):
        _set_cell_bg(cell, "1F497D")
        cell.paragraphs[0].clear()
        _add_run(cell.paragraphs[0], hdr,
                 bold=True, color=COLOR_WHITE, size_pt=10)

    row = table.add_row()
    for c in row.cells:
        _set_cell_bg(c, "F5F9FC")
    row.cells[0].paragraphs[0].clear()
    _add_run(row.cells[0].paragraphs[0], "{{ nfr_rows }}",
             italic=True, color=COLOR_MUTED, size_pt=9)


def _risk_table(doc):
    _heading(doc, "8. Risks", 1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    widths = [Cm(2.0), Cm(6.0), Cm(2.5), Cm(2.5), Cm(3.0)]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width

    hdr_cells = table.rows[0].cells
    for cell, hdr in zip(hdr_cells,
                         ["ID", "Description", "Likelihood", "Impact", "Mitigation"]):
        _set_cell_bg(cell, "1F497D")
        cell.paragraphs[0].clear()
        _add_run(cell.paragraphs[0], hdr,
                 bold=True, color=COLOR_WHITE, size_pt=10)

    row = table.add_row()
    for c in row.cells:
        _set_cell_bg(c, "F5F9FC")
    row.cells[0].paragraphs[0].clear()
    _add_run(row.cells[0].paragraphs[0], "{{ risk_rows }}",
             italic=True, color=COLOR_MUTED, size_pt=9)


def create_template(output_path: str = "templates/brd_template.docx"):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Cover Page ────────────────────────────────────────────────
    _cover_page(doc)

    # ── Section 1: Business Objectives ───────────────────────────
    _section_bullets(doc,
        "1. Business Objectives", "business_objectives_list",
        "Auto-filled from transcript")

    # ── Section 2: In Scope ───────────────────────────────────────
    _section_bullets(doc,
        "2. In Scope", "in_scope_list",
        "Auto-filled from transcript")

    # ── Section 3: Out of Scope ───────────────────────────────────
    _section_bullets(doc,
        "3. Out of Scope", "out_of_scope_list",
        "Auto-filled from transcript")

    # ── Section 4: Functional Requirements ───────────────────────
    _fr_table(doc)

    # ── Section 5: Non-Functional Requirements ────────────────────
    _nfr_table(doc)

    # ── Section 6: Assumptions ────────────────────────────────────
    _section_bullets(doc,
        "6. Assumptions", "assumptions_list",
        "Auto-filled from transcript")

    # ── Section 7: Constraints ────────────────────────────────────
    _section_bullets(doc,
        "7. Constraints", "constraints_list",
        "Auto-filled from transcript")

    # ── Section 8: Risks ──────────────────────────────────────────
    _risk_table(doc)

    # ── Attendees (after cover, before section 1) ─────────────────
    # Insert after cover page break — positioned in document flow
    _attendees_section(doc)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Template created: {output_path}")
    return output_path


if __name__ == "__main__":
    create_template()
