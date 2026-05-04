"""
BRD Builder — fills the company .docx template with BRDContent extracted by AI.

Strategy:
- Load template, find placeholder paragraphs by tag pattern {{ tag }}
- Replace inline placeholders (cover page fields) in-place
- Replace block placeholders (bullet lists, table rows) by rebuilding paragraphs/rows
- Never mutate the original template file
"""
import copy
import io
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from loguru import logger

from app.core.config import settings
from app.core.models import BRDContent, FunctionalRequirement, NonFunctionalRequirement, Risk


# ── Colors (must match template) ──────────────────────────────────────────────
_COLOR_TEXT    = RGBColor(0x26, 0x26, 0x26)
_COLOR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
_COLOR_HIGH    = RGBColor(0xC0, 0x00, 0x00)   # Red  — High risk/priority
_COLOR_MEDIUM  = RGBColor(0xED, 0x7D, 0x31)   # Orange — Medium
_COLOR_LOW     = RGBColor(0x70, 0xAD, 0x47)   # Green — Low

_PLACEHOLDER_RE = re.compile(r"\{\{\s*(\w+)\s*\}\}")


class BRDBuilderError(Exception):
    """Raised when template is missing, corrupt, or a required placeholder is absent."""


# ── Public API ─────────────────────────────────────────────────────────────────

def build_brd(content: BRDContent) -> bytes:
    """
    Fill the company BRD template with extracted content.
    Returns the filled .docx as raw bytes (ready for base64 encoding).

    Raises:
        BRDBuilderError: template not found or critical placeholder missing
    """
    template_path = Path(settings.brd_template_path)
    if not template_path.exists():
        raise BRDBuilderError(
            f"BRD template not found at '{template_path}'. "
            "Run scripts/create_template.py or set BRD_TEMPLATE_PATH in .env"
        )

    doc = Document(str(template_path))
    logger.info(f"Loaded template: {template_path}")

    _fill_inline_placeholders(doc, content)
    _fill_bullet_sections(doc, content)
    _fill_fr_table(doc, content)
    _fill_nfr_table(doc, content)
    _fill_risk_table(doc, content)

    buffer = io.BytesIO()
    doc.save(buffer)
    result = buffer.getvalue()

    logger.info(
        f"BRD built: {len(result)} bytes | "
        f"FR={len(content.functional_requirements)} | "
        f"NFR={len(content.non_functional_requirements)} | "
        f"Risks={len(content.risks)}"
    )
    return result


# ── Inline placeholder replacement ────────────────────────────────────────────

def _fill_inline_placeholders(doc: Document, content: BRDContent) -> None:
    """
    Replace {{ project_name }}, {{ meeting_date }} etc. in all paragraphs
    including inside table cells (cover page meta table).
    """
    replacements = {
        "project_name": content.project_name,
        "meeting_date": content.meeting_date,
    }

    # Body paragraphs
    for para in doc.paragraphs:
        _replace_inline(para, replacements)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_inline(para, replacements)


def _replace_inline(para, replacements: dict[str, str]) -> None:
    """Replace {{ key }} tokens within a paragraph's runs."""
    full_text = "".join(r.text for r in para.runs)
    match = _PLACEHOLDER_RE.search(full_text)
    if not match:
        return

    tag = match.group(1)
    if tag not in replacements:
        return

    new_text = _PLACEHOLDER_RE.sub(replacements[tag], full_text)

    # Clear all runs and put text in the first run to preserve formatting
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""


# ── Bullet list sections ───────────────────────────────────────────────────────

def _fill_bullet_sections(doc: Document, content: BRDContent) -> None:
    bullet_map = {
        "business_objectives_list": content.business_objectives,
        "in_scope_list":            content.in_scope,
        "out_of_scope_list":        content.out_of_scope,
        "assumptions_list":         content.assumptions,
        "constraints_list":         content.constraints,
        "attendees_list":           content.attendees,
    }

    for tag, items in bullet_map.items():
        _replace_bullet_placeholder(doc, tag, items)


def _replace_bullet_placeholder(doc: Document, tag: str, items: list[str]) -> None:
    """
    Find the paragraph containing {{ tag }}, remove it,
    and insert styled bullet paragraphs in its place.
    """
    placeholder_para = _find_paragraph(doc, tag)
    if placeholder_para is None:
        logger.warning(f"Placeholder '{{{{{tag}}}}}' not found in template — skipping")
        return

    parent = placeholder_para._element.getparent()
    insert_index = list(parent).index(placeholder_para._element)

    if not items:
        # Replace with "None identified" note
        items = ["None identified in transcript."]

    # Build new bullet paragraphs and insert before removing placeholder
    new_elements = []
    for item in items:
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), "1")
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)
        p.append(pPr)

        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")  # 11pt
        rPr.append(sz)
        r.append(rPr)

        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = item
        r.append(t)
        p.append(r)
        new_elements.append(p)

    # Insert new elements
    for offset, elem in enumerate(new_elements):
        parent.insert(insert_index + offset, elem)

    # Remove the original placeholder paragraph
    parent.remove(placeholder_para._element)
    logger.debug(f"Filled '{tag}' with {len(items)} bullet(s)")


# ── FR Table ───────────────────────────────────────────────────────────────────

def _fill_fr_table(doc: Document, content: BRDContent) -> None:
    table = _find_table_with_placeholder(doc, "fr_rows")
    if table is None:
        logger.warning("FR table placeholder not found — skipping")
        return

    # Remove placeholder row (last row)
    _remove_placeholder_row(table)

    for i, fr in enumerate(content.functional_requirements):
        row = table.add_row()
        bg = "F5F9FC" if i % 2 == 0 else "FFFFFF"
        _set_cell_bg(row.cells[0], bg)
        _set_cell_bg(row.cells[1], bg)
        _set_cell_bg(row.cells[2], bg)

        _write_cell(row.cells[0], fr.id, bold=True, size_pt=9)
        _write_cell(row.cells[1], fr.description, size_pt=10)
        _write_cell(
            row.cells[2], fr.priority,
            color=_priority_color(fr.priority), bold=True, size_pt=9
        )

    logger.debug(f"FR table filled: {len(content.functional_requirements)} rows")


# ── NFR Table ──────────────────────────────────────────────────────────────────

def _fill_nfr_table(doc: Document, content: BRDContent) -> None:
    table = _find_table_with_placeholder(doc, "nfr_rows")
    if table is None:
        logger.warning("NFR table placeholder not found — skipping")
        return

    _remove_placeholder_row(table)

    for i, nfr in enumerate(content.non_functional_requirements):
        row = table.add_row()
        bg = "F5F9FC" if i % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            _set_cell_bg(cell, bg)

        _write_cell(row.cells[0], nfr.id, bold=True, size_pt=9)
        _write_cell(row.cells[1], nfr.category, bold=True, size_pt=9,
                    color=_COLOR_TEXT)
        _write_cell(row.cells[2], nfr.description, size_pt=10)

    logger.debug(f"NFR table filled: {len(content.non_functional_requirements)} rows")


# ── Risk Table ─────────────────────────────────────────────────────────────────

def _fill_risk_table(doc: Document, content: BRDContent) -> None:
    table = _find_table_with_placeholder(doc, "risk_rows")
    if table is None:
        logger.warning("Risk table placeholder not found — skipping")
        return

    _remove_placeholder_row(table)

    for i, risk in enumerate(content.risks):
        row = table.add_row()
        bg = "F5F9FC" if i % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            _set_cell_bg(cell, bg)

        _write_cell(row.cells[0], risk.id, bold=True, size_pt=9)
        _write_cell(row.cells[1], risk.description, size_pt=10)
        _write_cell(row.cells[2], risk.likelihood.value,
                    color=_priority_color(risk.likelihood.value), bold=True, size_pt=9)
        _write_cell(row.cells[3], risk.impact.value,
                    color=_priority_color(risk.impact.value), bold=True, size_pt=9)
        _write_cell(row.cells[4], risk.mitigation, size_pt=9)

    logger.debug(f"Risk table filled: {len(content.risks)} rows")


# ── Helpers ────────────────────────────────────────────────────────────────────

def _find_paragraph(doc: Document, tag: str):
    """Find first paragraph containing {{ tag }} anywhere in document body."""
    pattern = re.compile(r"\{\{\s*" + re.escape(tag) + r"\s*\}\}")
    for para in doc.paragraphs:
        if pattern.search("".join(r.text for r in para.runs)):
            return para
    return None


def _find_table_with_placeholder(doc: Document, tag: str):
    """Find the table that contains a cell with {{ tag }}."""
    pattern = re.compile(r"\{\{\s*" + re.escape(tag) + r"\s*\}\}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if pattern.search("".join(r.text for r in para.runs)):
                        return table
    return None


def _remove_placeholder_row(table) -> None:
    """Remove the last row of the table (the placeholder row)."""
    tr = table.rows[-1]._tr
    tr.getparent().remove(tr)


def _write_cell(cell, text: str, bold=False, size_pt=10,
                color: RGBColor = None) -> None:
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color if color else _COLOR_TEXT


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _priority_color(value: str) -> RGBColor:
    mapping = {
        "High":   _COLOR_HIGH,
        "Medium": _COLOR_MEDIUM,
        "Low":    _COLOR_LOW,
    }
    return mapping.get(value, _COLOR_TEXT)
