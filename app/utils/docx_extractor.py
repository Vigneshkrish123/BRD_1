import io
from docx import Document
from loguru import logger


def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract plain text from a .docx transcript file.

    Handles:
    - Standard paragraphs
    - Tables (common in Teams/Zoom exported transcripts)
    - Preserves speaker: text structure where it exists in the doc
    """
    doc = Document(io.BytesIO(file_bytes))
    lines: list[str] = []

    # Extract from body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Extract from tables — Teams exports sometimes use table format
    for table in doc.tables:
        for row in table.rows:
            row_text = "  ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                lines.append(row_text)

    result = "\n".join(lines)
    logger.debug(f"DOCX extracted: {len(lines)} lines, {len(result)} chars")
    return result
